from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from docx.shared import RGBColor
from docx2pdf import convert
import os



from docx.shared import Pt

document=Document()
style = document.styles['Normal']
sections = document.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('This Resume has been created with Python')
italic_subhead2.italic=True
italic_subhead2.font.color.rgb = RGBColor(205,0,0)
subhead2.paragraph_format.left_indent = Inches(0.7)


#NAME
heading=document.add_heading('ANKITA MANDAL', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#Personal Details
subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('19.03.1996 | 9476486590 | ankitamandal193@gmail.com | Kharagpur, West Bengal, 721304')
italic_subhead2.italic=True
italic_subhead2.font.color.rgb = RGBColor(205,102,29)
subhead2.paragraph_format.left_indent = Inches(0.7)

#Experience
subhead=document.add_heading('Experience', 2)
subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Vinod Gupta School of Management, IIT Kharagpur | October 2019 - Present')
italic_subhead2.font.color.rgb = RGBColor(0,100,245)
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.7)

paragraph = document.add_paragraph('Assist with the planning and preparatory work of the Office’s work programme '
                                   'and/or project initiatives. Monitor status of programme and/or project proposals '
                                   'and receipt of documentation for review and approval, verifying that information '
                                   'is in compliance with applicable rules, regulations, policies, procedures and '
                                   'guidelines. Monitor the status of programme and/or project outcomes and '
                                   'deliverables and inform the supervisor of any discrepancies.',
                  style='List Number 2')
font =paragraph.style.font
font.name='Arial'
font.size=Pt(10)
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Compile, summarize and present a variety of information and data to the '
                                   'supervisor on issues pertinent to the Office’s work programme..',
                  style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Input complete data and process administrative actions on programme and/or '
                                   'project delivery in the enterprise resource planning (ERP) system. Based on '
                                   'information in the enterprise resource planning (ERP) system, inform the '
                                   'supervisor of inconsistencies and shortfalls in delivery and status of '
                                   'allocations. Distribute project documents to concerned parties upon approval of '
                                   'supervisor.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Provide administrative support for the organization of seminars, workshops, '
                                   'meetings and other events.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Maintain and update databases. Perform basic searches for information and prepare '
                                   'and update periodic reports, background information, briefing notes and '
                                   'statistical summaries.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Respond to requests for general information on programme and/or project related '
                                   'matters.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Keep abreast of changes to relevant programme-related policies, procedures, '
                                   'guidelines and processes and share information with concerned parties, '
                                   'providing further clarification as required', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Provide general office management support, including attending meetings and '
                                   'drafting correspondence.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

paragraph = document.add_paragraph('Perform other relevant duties as assigned.', style='List Number 2')
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.paragraph_format.line_spacing = Inches(0.3)

#Portfolio
subhead=document.add_heading('Portfolio', 2)
subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('1: Gender Classification based on movie dialogue')
italic_subhead2.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
subhead2.paragraph_format.left_indent = Inches(0.5)



paragraph = document.add_paragraph('''Solved by ''',style='List Bullet 2')
paragraph.add_run('''classification model.''',).bold=True
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('Prepare the data',style='List Bullet 2')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Created ''',style='List Bullet 2')
paragraph.add_run('''training and test model and performed the test.''',).bold=True
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Created ''',style='List Bullet 2')
paragraph.add_run('''Pipeline and confusion matrix''',).bold=True
paragraph.paragraph_format.left_indent = Inches(1)


paragraph = document.add_paragraph('''Performed and checked accuracy of the model of ''',style='List Bullet 2')
paragraph.add_run('''Logistic regression''',).bold=True
paragraph.paragraph_format.left_indent = Inches(1)


paragraph = document.add_paragraph('It was 0.7 ',style='List Bullet 2')
paragraph.paragraph_format.left_indent = Inches(1)
font =paragraph.style.font
font.name='Calibri'
font.size=Pt(12)

subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Git-hub URL:')
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.5)

link = document.add_paragraph('')
italic_link = link.add_run('https://github.com/ankitavgsom/personal/blob/main/gender%20classification%20based%20on%20movie%20dialogue.ipynb')
italic_link.italic=True
italic_link.underline=True
italic_link.font.color.rgb = RGBColor(0,0,139)
link.paragraph_format.left_indent = Inches(0.5)



subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('2: Relationship between Diabetes, Diabetes Distress and Emotional Burden in Malaysian Population: A Case Study')
italic_subhead2.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
subhead2.paragraph_format.left_indent = Inches(0.5)

paragraph = document.add_paragraph('''Used ''',style='List Bullet 2')
paragraph.add_run('''Statistical Methodology and Interpretation''').bold=True
paragraph.add_run(''' in order to uncover the pattern and trends of the data''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Designed ''',style='List Bullet 2')
paragraph.add_run('''Hypothesis Testing''').bold=True
paragraph.add_run(''' for making determination related to the population''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Designed ''',style='List Bullet 2')
paragraph.add_run('''Linear Regression''').bold=True
paragraph.add_run(''' for modelling the relationship between dependent and independent variable''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Formed ''',style='List Bullet 2')
paragraph.add_run('''Ordinal Regression''').bold=True
paragraph.add_run(''' to predict the dependent variable with 'ordered' multiple categories and independent variables''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Interpreted ''',style='List Bullet 2')
paragraph.add_run('''Model Fitting Information''').bold=True
paragraph.add_run(''' for testing goodness-of-fit''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Analyzed ''',style='List Bullet 2')
paragraph.add_run('''Pseudo R-Square''').bold=True
paragraph.add_run(''' to summarize the proportion of variance in the dependent variable associated with the predictor (independent) variables''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''Created ''',style='List Bullet 2')
paragraph.add_run('''Dummy variable''').bold=True
paragraph.add_run(''' to represent the subgroups in the data''')
paragraph.paragraph_format.left_indent = Inches(1)

subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Project URL:')
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.5)

link = document.add_paragraph('')
italic_link = link.add_run('https://drive.google.com/file/d/0B5xMV5g_wrGaRG5uR1dRdTlYdm9hZmotc09UT3B4cHdoN3o0/view?usp=sharing&resourcekey=0-xfc-HnHwQoiqPe8SZC6JyQ')
italic_link.italic=True
italic_link.underline=True
italic_link.font.color.rgb = RGBColor(0,0,139)
link.paragraph_format.left_indent = Inches(0.5)

subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('3: Build a Resume with Python')
italic_subhead2.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
subhead2.paragraph_format.left_indent = Inches(0.5)




paragraph = document.add_paragraph('''Created ''',style='List Bullet 2')
paragraph.add_run('''a Word File''').bold=True
paragraph.add_run(''' with Python''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''The word file has been''',style='List Bullet 2')
paragraph.add_run(''' formatted''').bold=True
paragraph.add_run(''' with Python''')
paragraph.paragraph_format.left_indent = Inches(1)

paragraph = document.add_paragraph('''The word file has been''',style='List Bullet 2')
paragraph.add_run(''' converted to pdf''').bold=True
paragraph.add_run(''' with Python''')
paragraph.paragraph_format.left_indent = Inches(1)

subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Git-hub URL:')
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.5)

link = document.add_paragraph('')
italic_link = link.add_run('https://github.com/ankitavgsom/personal/blob/main/cv.py')
italic_link.italic=True
italic_link.underline=True
italic_link.font.color.rgb = RGBColor(0,0,139)
link.paragraph_format.left_indent = Inches(0.5)



font =paragraph.style.font
font.name='Calibri'
font.size=Pt(11)


#Education
subhead=document.add_heading('Education', 2)
subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Master in Applied Statistics and Informatics | 2017-2019')
italic_subhead2.font.color.rgb = RGBColor(0,100,245)
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.7)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Central University of Orissa | Odisha')
subhead2.paragraph_format.left_indent = Inches(1)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Course Work: Statistics, Probability, Inference, Regression and Programming with R, SPSS, MATLAB, C')
italic_subhead2.bold=True
subhead2.paragraph_format.left_indent = Inches(1)




subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Bachelor in Mathematics (H) | 2013-2017')
italic_subhead2.font.color.rgb = RGBColor(0,100,245)
italic_subhead2.italic=True
subhead2.paragraph_format.left_indent = Inches(0.7)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Vidyasagar University | West Bengal')
subhead2.paragraph_format.left_indent = Inches(1)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Course Work: Mathematics, Physics, Chemistry')
italic_subhead2.bold=True
subhead2.paragraph_format.left_indent = Inches(1)

#Certification
subhead=document.add_heading('Computer Skills', 2)

subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Tools')
italic_subhead2.bold=True
subhead2.paragraph_format.left_indent = Inches(0.5)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('R, PYTHON, SQL, C, SPSS, MATLAB, MS OFFICE, Numpy, Jupyter, Scikit-Learn, Tableau')
subhead2.paragraph_format.left_indent = Inches(1)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Techniques')
italic_subhead2.bold=True
subhead2.paragraph_format.left_indent = Inches(0.5)


subhead2=document.add_paragraph('')
italic_subhead2 = subhead2.add_run('Statistical Methods, Statistical Modelling, Data Analysis, Interpretation and Visualization, Sampling, ANOVA Test, Machine Learning Skills')
subhead2.paragraph_format.left_indent = Inches(1)



document.save('cv.docx')
os.startfile('cv.docx')

convert("cv.docx", "CV_Pthon_Ankita Mandal.pdf")
